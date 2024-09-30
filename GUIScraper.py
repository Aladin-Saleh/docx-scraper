import os
import docx
import re
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext

class KeywordSearchApp:
    def __init__(self, master):
        self.master = master
        master.title("Recherche de Mots Clés dans des fichiers .docx")

        self.label = tk.Label(master, text="Entrez les mots-clés (séparés par des virgules) :")
        self.label.pack()

        self.keywords_entry = tk.Entry(master, width=50)
        self.keywords_entry.pack()

        self.select_folder_button = tk.Button(master, text="Sélectionner un dossier", command=self.select_folder)
        self.select_folder_button.pack()

        self.results_text = scrolledtext.ScrolledText(master, width=80, height=20)
        self.results_text.pack()

        self.search_button = tk.Button(master, text="Lancer la recherche", command=self.process_folder)
        self.search_button.pack()

        self.folder_path = ""

    def select_folder(self):
        self.folder_path = filedialog.askdirectory()
        if not self.folder_path:
            messagebox.showwarning("Avertissement", "Veuillez sélectionner un dossier.")

    def load_docx(self, file_path):
        return docx.Document(file_path)

    def extract_text(self, doc):
        text = '\n'.join([para.text for para in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += '\n' + cell.text
        return text

    def find_emails(self, text):
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        return re.findall(email_pattern, text)

    def find_phone_numbers(self, text):
        phone_pattern = r'\+?\d[\d -]{8,12}\d'
        return re.findall(phone_pattern, text)

    def find_keywords(self, text, keywords):
        found_keywords = {}
        for keyword in keywords:
            pattern = rf'\b{keyword}\b'
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                found_keywords[keyword] = len(matches)
        return found_keywords

    def process_folder(self):
        keywords = [keyword.strip() for keyword in self.keywords_entry.get().split(',')]
        if not self.folder_path or not keywords:
            messagebox.showwarning("Avertissement", "Veuillez sélectionner un dossier et entrer des mots-clés.")
            return

        results = {}
        for filename in os.listdir(self.folder_path):
            if filename.endswith(".docx"):
                doc = self.load_docx(os.path.join(self.folder_path, filename))
                text = self.extract_text(doc)
                emails = self.find_emails(text)
                phones = self.find_phone_numbers(text)
                found_keywords = self.find_keywords(text, keywords)

                if found_keywords:
                    total_keywords = len(found_keywords)  
                    results[filename] = {
                        'emails': emails,
                        'phones': phones,
                        'found_keywords': found_keywords,
                        'total_keywords': total_keywords
                    }

        sorted_results = sorted(results.items(), key=lambda item: item[1]['total_keywords'], reverse=True)

        self.results_text.delete(1.0, tk.END)
        for filename, data in sorted_results:
            self.results_text.insert(tk.END, f"Fichier : {filename}\n")
            self.results_text.insert(tk.END, f"Emails trouvés : {data['emails']}\n")
            self.results_text.insert(tk.END, f"Téléphones trouvés : {data['phones']}\n")
            self.results_text.insert(tk.END, f"Mots clés trouvés : {data['found_keywords']}\n\n")

root = tk.Tk()
app = KeywordSearchApp(root)
root.mainloop()
